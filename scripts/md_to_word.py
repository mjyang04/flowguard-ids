#!/usr/bin/env python3
"""Convert thesis_draft_en.md to a Word document, preserving structure and formatting."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, bold=False, italic=False, size=None, color=None, monospace=False):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if monospace:
        run.font.name = "Courier New"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rPr.insert(0, rFonts)
    else:
        run.font.name = "Times New Roman"


def add_paragraph(doc, style="Normal", text="", bold=False, italic=False,
                  size=None, align=None, monospace=False, keep_with_next=False):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if text:
        run = p.add_run(text)
        run.font.name = "Courier New" if monospace else "Times New Roman"
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size or 12)
    return p


def parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """Parse inline formatting. Returns list of (content, bold, italic)."""
    segments = []
    # Match **bold**, *italic*, `code`
    pattern = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end:m.start()], False, False))
        if m.group(1):  # **bold**
            segments.append((m.group(1), True, False))
        elif m.group(2):  # *italic*
            segments.append((m.group(2), False, True))
        elif m.group(3):  # `code`
            segments.append((m.group(3), False, False))
        last_end = m.end()
    if last_end < len(text):
        segments.append((text[last_end:], False, False))
    return segments


def add_inline_paragraph(doc, text: str, style="Normal", size=None,
                         align=None, keep_with_next=False,
                         bold=False, italic=False, monospace=False):
    p = doc.add_paragraph(style=style)
    # Default to justified for body text
    p.alignment = align if align else WD_ALIGN_PARAGRAPH.JUSTIFY
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    segments = parse_inline(text)
    if not segments:
        return p
    for content, b, it in segments:
        run = p.add_run(content)
        use_bold = bold or b
        use_italic = italic or it
        run.font.name = "Courier New" if monospace else "Times New Roman"
        run.bold = use_bold
        run.italic = use_italic
        run.font.size = Pt(size or 12)
    return p


def add_heading(doc, text: str, level: int):
    styles = {
        1: "Title",
        2: "Heading 1",
        3: "Heading 2",
        4: "Heading 3",
        5: "Heading 4",
        6: "Heading 5",
    }
    style = styles.get(level, "Heading 6")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14) if level <= 2 else Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.keep_with_next = True
    return p


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse markdown table. Returns (headers, rows)."""
    # Find table start (has | at start and end)
    header_line = None
    for i, line in enumerate(lines):
        if re.match(r"^\|.+\|$", line.strip()):
            header_line = i
            break
    if header_line is None:
        return [], []

    headers = [h.strip() for h in lines[header_line].strip("|").split("|")]
    # Skip separator line
    data_lines = lines[header_line + 2:]
    rows = []
    for line in data_lines:
        if re.match(r"^\|.+\|$", line.strip()):
            row = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(row)
    return headers, rows


def add_table(doc, headers: list[str], rows: list[list[str]]):
    if not headers or not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell in enumerate(row):
            if c_idx < len(cells):
                cells[c_idx].text = cell
                for p in cells[c_idx].paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
    return table


def convert_md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()

    # Document-level defaults
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Heading styles
    for lvl in range(1, 7):
        try:
            s = doc.styles[f"Heading {lvl}"]
            s.font.name = "Times New Roman"
            s.font.size = Pt(14 if lvl <= 2 else 12)
            s.font.bold = True
        except KeyError:
            pass
    try:
        title_style = doc.styles["Title"]
        title_style.font.name = "Times New Roman"
        title_style.font.size = Pt(14)
        title_style.font.bold = True
    except KeyError:
        pass

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_abstract = False
    abstract_paragraphs = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip table of contents marker
        if "<!-- TOC" in line or "<!-- toc" in line.lower():
            i += 1
            continue

        # Heading 1: Title or major section
        if re.match(r"^# ", line):
            text = line[2:].strip()
            add_heading(doc, text, 1)

        elif re.match(r"^## ", line):
            text = line[3:].strip()
            add_heading(doc, text, 2)

        elif re.match(r"^### ", line):
            text = line[4:].strip()
            add_heading(doc, text, 3)

        elif re.match(r"^#### ", line):
            text = line[5:].strip()
            add_heading(doc, text, 4)

        elif re.match(r"^##### ", line):
            text = line[6:].strip()
            add_heading(doc, text, 5)

        elif re.match(r"^###### ", line):
            text = line[7:].strip()
            add_heading(doc, text, 6)

        # Horizontal rule
        elif re.match(r"^---+$", line) or re.match(r"^\*\*\*+$", line):
            pass  # Word handles section breaks via headings

        # Table: detect consecutive table lines
        elif re.match(r"^\|.+\|$", line.strip()) and line.strip():
            table_lines = []
            while i < len(lines) and re.match(r"^\|.+\|$", lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_table(table_lines)
            if headers:
                add_table(doc, headers, rows)
            doc.add_paragraph()  # spacing after table
            continue

        # Blockquote / abstract marker
        elif line.startswith("> "):
            text = line[2:].strip()
            add_inline_paragraph(doc, text, size=12, italic=True)
            # Check next non-empty line for more abstract
            j = i + 1
            while j < len(lines) and lines[j].startswith("> "):
                add_inline_paragraph(doc, lines[j][2:].strip(), size=12, italic=True)
                j += 1
                i = j
            continue

        # Bullet list
        elif re.match(r"^[-*] ", line):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            segments = parse_inline(text)
            for content, bold, italic in segments:
                run = p.add_run(content)
                set_font(run, bold=bold, italic=italic, size=12)
            i += 1
            continue

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line).strip()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            segments = parse_inline(text)
            for content, bold, italic in segments:
                run = p.add_run(content)
                set_font(run, bold=bold, italic=italic, size=12)
            i += 1
            continue

        # Code block
        elif line.strip().startswith("```"):
            # Collect code block
            lang = line.strip()[3:]
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            i += 1  # skip closing ```
            continue

        # Math block ($$...$$) — store as plain text
        elif re.match(r"^\$\$", line):
            math_lines = []
            i += 1
            while i < len(lines) and not re.match(r"^\$\$", lines[i]):
                math_lines.append(lines[i])
                i += 1
            math_text = " ".join(math_lines)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(math_text)
            run.italic = True
            run.font.size = Pt(11)
            i += 1
            continue

        # Inline math ($...$) — convert to plain text with warning
        elif "$" in line:
            # Replace inline math $...$ with plain text
            new_line = re.sub(r"\$([^$]+)\$", r"\1", line)
            if new_line.strip():
                add_inline_paragraph(doc, new_line.strip(), size=12)
            i += 1
            continue

        # Empty line
        elif not line.strip():
            i += 1
            continue

        # Regular paragraph
        else:
            text = line
            # Check if next line continues (not empty, not a heading, not a list, not code)
            j = i + 1
            while (j < len(lines) and lines[j].strip() and
                   not re.match(r"^(#{1,6} |[-*] |\d+\. |```|\$\$)", lines[j]) and
                   not re.match(r"^\|.+\|$", lines[j].strip())):
                text += " " + lines[j].strip()
                j += 1
            add_inline_paragraph(doc, text, size=12)
            i = j
            continue

        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    docs_dir = Path(__file__).parent.parent / "docs"
    for md_name, docx_name in [
        ("thesis_draft_en.md", "thesis_draft_en.docx"),
        ("thesis_draft.md", "thesis_draft.docx"),
    ]:
        md_path = docs_dir / md_name
        docx_path = docs_dir / docx_name
        if md_path.exists():
            convert_md_to_docx(md_path, docx_path)
        else:
            print(f"Skipped (not found): {md_path}")
